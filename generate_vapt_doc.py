from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Color palette ──
DARK_BLUE  = RGBColor(0x00, 0x2B, 0x5B)
MID_BLUE   = RGBColor(0x00, 0x57, 0xB8)
LIGHT_BLUE = RGBColor(0xEB, 0xF3, 0xFF)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY_TEXT  = RGBColor(0x44, 0x44, 0x44)
RED_BG     = RGBColor(0xFF, 0xF0, 0xF0)
RED_BORDER = RGBColor(0xDC, 0x26, 0x26)
AMBER_BG   = RGBColor(0xFF, 0xF8, 0xE6)
GREEN_BG   = RGBColor(0xD1, 0xFA, 0xE5)
HEADER_BG  = RGBColor(0x00, 0x2B, 0x5B)


# ════════════════════════════════════════════════════
# Helper utilities
# ════════════════════════════════════════════════════

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        if color:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   'single')
            el.set(qn('w:sz'),    '18')
            el.set(qn('w:color'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
            tcBorders.append(el)
    tcPr.append(tcBorders)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(16 if level == 1 else 13)
    if level == 1:
        # bottom border via paragraph border
        pPr   = p._p.get_or_add_pPr()
        pBdr  = OxmlElement('w:pBdr')
        bot   = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '8')
        bot.set(qn('w:color'), '0057B8')
        pBdr.append(bot)
        pPr.append(pBdr)
    return p

def section_label(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size  = Pt(8)
    run.font.bold  = True
    run.font.color.rgb = MID_BLUE
    p.paragraph_format.space_after = Pt(0)
    return p

def body(doc, text):
    p   = doc.add_paragraph(text)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = GREY_TEXT
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(doc, text, indent=0):
    p   = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = GREY_TEXT
    p.paragraph_format.left_indent   = Cm(0.5 + indent * 0.4)
    p.paragraph_format.space_after   = Pt(3)
    return p

def numbered(doc, text):
    p   = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = GREY_TEXT
    p.paragraph_format.space_after   = Pt(3)
    return p

def make_table(doc, headers, rows, col_widths=None, stripe=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, HEADER_BG)
        p   = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = LIGHT_BLUE if (stripe and ri % 2 == 0) else WHITE
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p   = cell.paragraphs[0]
            if isinstance(val, tuple):   # (text, bold)
                run = p.add_run(val[0])
                run.bold = val[1]
            else:
                run = p.add_run(val)
            run.font.size = Pt(9.5)
            run.font.color.rgb = GREY_TEXT

    if col_widths:
        set_col_widths(table, col_widths)

    doc.add_paragraph()
    return table

def info_box(doc, label, text, kind='info'):
    bg  = {
        'info':    LIGHT_BLUE,
        'warning': AMBER_BG,
        'danger':  RED_BG,
    }.get(kind, LIGHT_BLUE)

    table = doc.add_table(rows=1, cols=1)
    cell  = table.rows[0].cells[0]
    set_cell_bg(cell, bg)

    border_color = {
        'info':    MID_BLUE,
        'warning': RGBColor(0xF5, 0x9E, 0x0B),
        'danger':  RED_BORDER,
    }.get(kind, MID_BLUE)
    set_cell_border(cell, left=border_color)

    p = cell.paragraphs[0]
    if label:
        bold_run = p.add_run(label + '\n')
        bold_run.bold = True
        bold_run.font.size = Pt(9.5)
    text_run = p.add_run(text)
    text_run.font.size = Pt(9.5)
    text_run.font.color.rgb = GREY_TEXT

    doc.add_paragraph()

def page_break(doc):
    doc.add_page_break()

def sub_heading(doc, text):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(11)
    return p

def inline_code(run_or_para, text):
    """Add monospace inline text to an existing paragraph."""
    if isinstance(run_or_para, str):
        return f"`{text}`"
    run = run_or_para.add_run(f' {text} ')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return run


# ════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run('CONFIDENTIAL  ·  SECURITY ASSESSMENT')
run.font.size  = Pt(9)
run.font.color.rgb = MID_BLUE
run.font.bold  = True

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run('Vulnerability Assessment & Penetration Testing')
run.font.size  = Pt(28)
run.font.bold  = True
run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(40)
run = p.add_run('Scope Document')
run.font.size  = Pt(22)
run.font.color.rgb = MID_BLUE

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run('MIT Billing Application — Paytm POS Deployment')
run.font.size  = Pt(14)
run.font.color.rgb = GREY_TEXT
run.font.bold  = True

doc.add_paragraph()

meta = [
    ('Document Version', '1.0'),
    ('Date', 'March 2026'),
    ('Classification', 'Confidential – For VAPT Team Only'),
    ('Testing Type', 'Black-box + Grey-box VAPT'),
    ('Application Name', 'MIT Billing System'),
    ('Platform Target', 'Paytm POS (Android) + Web Admin + REST API'),
]
table = doc.add_table(rows=len(meta), cols=2)
table.style = 'Table Grid'
for i, (lbl, val) in enumerate(meta):
    row  = table.rows[i]
    bg   = LIGHT_BLUE if i % 2 == 0 else WHITE
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], bg)
    r0 = row.cells[0].paragraphs[0].add_run(lbl)
    r0.bold = True
    r0.font.size = Pt(9.5)
    r0.font.color.rgb = DARK_BLUE
    r1 = row.cells[1].paragraphs[0].add_run(val)
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = GREY_TEXT
set_col_widths(table, [5.5, 10])

page_break(doc)


# ════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════
section_label(doc, 'Navigation')
heading(doc, 'Table of Contents')
toc = [
    '1.  Executive Summary',
    '2.  Application Overview',
    '3.  System Architecture',
    '4.  Technology Stack',
    '5.  Components in Scope',
    '6.  API Endpoints Reference',
    '7.  Authentication & Authorization Flow',
    '8.  Payment Flows',
    '9.  Data Sensitivity & Classification',
    '10. VAPT Scope Definition',
    '11. Testing Focus Areas & Test Cases',
    '12. Out-of-Scope Items',
    '13. Testing Environment Details',
    '14. Engagement Rules & Contact',
]
for item in toc:
    p   = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.color.rgb = GREY_TEXT
    p.paragraph_format.space_after = Pt(4)
page_break(doc)


# ════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════
section_label(doc, 'Section 1')
heading(doc, '1.  Executive Summary')

body(doc, ('This document defines the scope, objectives, and technical details for the '
           'Vulnerability Assessment and Penetration Testing (VAPT) of the MIT Billing System '
           'prior to its commercial launch on the Paytm POS (Point of Sale) platform.'))

body(doc, 'The MIT Billing System is a multi-component payment and billing platform consisting of:')
bullet(doc, 'A Flutter Android client application deployed on Paytm EDC (Electronic Data Capture) POS terminals')
bullet(doc, 'A FastAPI REST backend providing all business logic and data persistence')
bullet(doc, 'A React web-based Admin Panel for merchant / admin management')

body(doc, ('The application handles real financial transactions (UPI, Card via Paytm EDC, and Cash) '
           'and stores sensitive payment data, making a thorough security assessment critical before '
           'production launch.'))

info_box(doc,
    '⚠  Purpose of this Document',
    ('This document is prepared to brief the VAPT team on the application\'s architecture, '
     'components, endpoints, authentication mechanisms, and data flows so that an effective '
     'and comprehensive security test can be conducted. All testing must be performed on the '
     'designated staging/test environment only.'),
    kind='warning')

page_break(doc)


# ════════════════════════════════════════════════════
# 2. APPLICATION OVERVIEW
# ════════════════════════════════════════════════════
section_label(doc, 'Section 2')
heading(doc, '2.  Application Overview')

sub_heading(doc, 'What the Application Does')
body(doc, ('MIT Billing System is a billing and payment management platform designed for institutions '
           '(e.g., colleges, hostels, service centers) to accept and track payments from customers '
           'at physical POS counters.'))

make_table(doc,
    ['User Type', 'Interface', 'Primary Actions'],
    [
        ('Admin / Merchant', 'React Web Admin Panel',
         'Manage machines, view dashboards, configure services, manage locations, export reports, view alerts'),
        ('Machine / Cashier', 'Flutter Android App (Paytm POS)',
         'Log in as a machine, select services, collect payments (UPI / Card / Cash), print bills, sync transactions'),
    ],
    col_widths=[3.5, 4.5, 8.5],
)

sub_heading(doc, 'Business Context')
for b in [
    'Each physical POS device is registered as a Machine in the system with unique credentials.',
    'Machines belong to Locations (e.g., "Counter A – Block 1").',
    'Services (items for billing) are configured per-machine from the admin panel.',
    'Payments are collected via Paytm EDC SDK (card/UPI), UPI QR, or Cash.',
    'Transactions sync to the central backend in real time or queue locally when offline.',
    'Bills are printed via the Paytm POS\'s built-in thermal printer.',
]:
    bullet(doc, b)

sub_heading(doc, 'Application ID / Package Name')
make_table(doc,
    ['Platform', 'Identifier'],
    [
        ('Android (Paytm POS)', 'com.mit.billing'),
        ('Backend API',         'Billing Admin API v1.0.0'),
    ],
    col_widths=[6, 10.5],
)
page_break(doc)


# ════════════════════════════════════════════════════
# 3. SYSTEM ARCHITECTURE
# ════════════════════════════════════════════════════
section_label(doc, 'Section 3')
heading(doc, '3.  System Architecture')

sub_heading(doc, 'Component Overview')
make_table(doc,
    ['Component', 'Technology', 'Role'],
    [
        ('Flutter App (POS Terminal)', 'Flutter / Dart — Android', 'Cashier-facing billing UI on Paytm POS device'),
        ('FastAPI Backend', 'Python / FastAPI / SQLAlchemy', 'Central REST API; business logic & data persistence'),
        ('PostgreSQL Database', 'PostgreSQL (SQLite for local dev)', 'Persistent storage for payments, machines, users, logs'),
        ('React Admin Panel', 'React + TypeScript + Vite', 'Web dashboard for admin/merchant management'),
    ],
    col_widths=[4.5, 5.5, 6.5],
)

sub_heading(doc, 'Communication Channels')
make_table(doc,
    ['From', 'To', 'Protocol', 'Auth Method'],
    [
        ('Flutter App (POS)', 'FastAPI Backend', 'HTTPS / REST', 'JWT Bearer Token'),
        ('React Admin', 'FastAPI Backend', 'HTTPS / REST', 'JWT Bearer Token'),
        ('Flutter App', 'Paytm EDC SDK', 'Local IPC (Android Intent / SDK)', 'Paytm SDK credentials'),
        ('Paytm EDC SDK', 'Paytm Payment Gateway', 'HTTPS (managed by SDK)', 'Paytm merchant credentials'),
    ],
    col_widths=[4, 4.5, 5, 4],
)

sub_heading(doc, 'Deployment Notes')
for b in [
    'Backend is hosted on a cloud server (production). A staging instance will be provided for VAPT.',
    'The Flutter app is side-loaded on Paytm POS terminals (APK deployment).',
    'Admin panel is a static web app served over HTTPS.',
    'CORS is configured on the backend to allow specific origins only.',
]:
    bullet(doc, b)

page_break(doc)


# ════════════════════════════════════════════════════
# 4. TECHNOLOGY STACK
# ════════════════════════════════════════════════════
section_label(doc, 'Section 4')
heading(doc, '4.  Technology Stack')

make_table(doc,
    ['Layer', 'Technology', 'Notes'],
    [
        ('Mobile App', 'Flutter (Dart)', 'Deployed on Paytm Android POS'),
        ('', 'Riverpod', 'State management'),
        ('', 'Dio', 'HTTP client with interceptors'),
        ('', 'GoRouter', 'App navigation'),
        ('', 'SharedPreferences', 'Local storage (offline queue, settings)'),
        ('', 'Paytm EDC SDK', 'Card & UPI payment processing'),
        ('Backend', 'Python / FastAPI', 'Async REST API'),
        ('', 'SQLAlchemy', 'ORM'),
        ('', 'PostgreSQL', 'Production database'),
        ('', 'JWT (HS256)', 'Access + Refresh token auth'),
        ('', 'Bcrypt', 'Password hashing'),
        ('', 'Pydantic v2', 'Request/response validation'),
        ('Admin Panel', 'React + TypeScript', 'Vite build tool'),
        ('', 'TanStack Query', 'API state & caching'),
        ('', 'shadcn/ui', 'UI component library'),
        ('', 'Recharts', 'Analytics charts'),
    ],
    col_widths=[3.5, 5, 8],
)
page_break(doc)


# ════════════════════════════════════════════════════
# 5. COMPONENTS IN SCOPE
# ════════════════════════════════════════════════════
section_label(doc, 'Section 5')
heading(doc, '5.  Components in Scope')

make_table(doc,
    ['#', 'Component', 'Type', 'Testing Priority'],
    [
        ('1', 'FastAPI REST Backend API', 'Web Application / API', '🔴 Critical'),
        ('2', 'Flutter Android Application (Paytm POS)', 'Mobile Application (Android)', '🔴 Critical'),
        ('3', 'React Admin Web Panel', 'Web Application', '🟠 High'),
        ('4', 'JWT Authentication Implementation', 'Authentication Mechanism', '🔴 Critical'),
        ('5', 'Offline Sync Queue (SharedPreferences)', 'Mobile Local Storage', '🟠 High'),
        ('6', 'Payment Submission Flow', 'Business Logic / API', '🔴 Critical'),
        ('7', 'Paytm EDC SDK Integration (app-layer only)', 'SDK Interface', '🟡 Medium'),
    ],
    col_widths=[1, 6, 4.5, 3],
)

info_box(doc,
    '📌 Note on Paytm EDC SDK',
    ('The Paytm EDC SDK itself (payment gateway communication) is managed by Paytm and is NOT in scope. '
     'However, how our application interfaces with the SDK — including how transaction results are handled '
     'and recorded — IS in scope.'),
    kind='info')

page_break(doc)


# ════════════════════════════════════════════════════
# 6. API ENDPOINTS REFERENCE
# ════════════════════════════════════════════════════
section_label(doc, 'Section 6')
heading(doc, '6.  API Endpoints Reference')

body(doc, 'All API endpoints are prefixed with /v1. Base URL to be provided separately for the staging environment.')

sub_heading(doc, 'Authentication Endpoints')
make_table(doc,
    ['Method', 'Endpoint', 'Auth Required', 'Description'],
    [
        ('POST', '/v1/auth/login',          'No',            'Admin/user login → returns JWT access + refresh token'),
        ('POST', '/v1/auth/machine-login',  'No',            'Machine (POS device) login → returns machine JWT'),
        ('POST', '/v1/auth/refresh',        'Refresh Token', 'Exchange refresh token for new access token'),
        ('POST', '/v1/auth/logout',         'Yes',           'Logout (client-side token invalidation)'),
        ('GET',  '/v1/auth/me',             'Yes',           'Get current user/machine info; updates machine status to "online"'),
    ],
    col_widths=[2, 5, 3.5, 6],
)

sub_heading(doc, 'Machine Management Endpoints')
make_table(doc,
    ['Method', 'Endpoint', 'Auth Required', 'Description'],
    [
        ('GET',    '/v1/machines/',     'Admin JWT', 'List all machines'),
        ('POST',   '/v1/machines/',     'Admin JWT', 'Create a new machine (auto-generates username)'),
        ('GET',    '/v1/machines/{id}', 'Admin JWT', 'Get machine details by ID'),
        ('PUT',    '/v1/machines/{id}', 'Admin JWT', 'Update machine (name, location, UPI ID, password)'),
        ('DELETE', '/v1/machines/{id}', 'Admin JWT', 'Delete a machine'),
    ],
    col_widths=[2, 5, 3.5, 6],
)

sub_heading(doc, 'Payment Endpoints')
make_table(doc,
    ['Method', 'Endpoint', 'Auth Required', 'Description'],
    [
        ('GET',  '/v1/payments',                      'JWT', 'Get all payments (filters: period, method, status, date range, machine_id, pagination)'),
        ('GET',  '/v1/machines/{machine_id}/payments','JWT', 'Get payments for a specific machine with filters'),
        ('POST', '/v1/payments',                      'JWT', 'Create a payment record (manual entry)'),
        ('GET',  '/v1/payments/{payment_id}',         'JWT', 'Get a specific payment'),
    ],
    col_widths=[2, 6, 3, 5.5],
)

sub_heading(doc, 'Sync Endpoints')
make_table(doc,
    ['Method', 'Endpoint', 'Auth Required', 'Description'],
    [
        ('POST', '/v1/sync/push',               'Machine JWT', 'Push offline-queued payments from POS to server'),
        ('POST', '/v1/sync/pull',               'Machine JWT', 'Pull latest services & config from server to POS'),
        ('GET',  '/v1/sync/status/{machine_id}','JWT',         'Get sync status of a machine'),
    ],
    col_widths=[2, 5, 3.5, 6],
)

sub_heading(doc, 'Other Endpoints')
make_table(doc,
    ['Module', 'Prefix', 'Examples'],
    [
        ('Dashboard',  '/v1/dashboard',  'Analytics summary, revenue charts'),
        ('Services',   '/v1/services',   'CRUD for billable services per machine'),
        ('Logs',       '/v1/logs',       'Activity logs, export logs to CSV'),
        ('Analytics',  '/v1/analytics',  'Revenue analytics, uptime stats, export to CSV'),
        ('Alerts',     '/v1/alerts',     'Machine status alerts (offline, maintenance)'),
        ('Config',     '/v1/config',     'Bill config (org name, UPI ID, bill prefix)'),
        ('Locations',  '/v1/locations',  'Location management (CRUD)'),
        ('Health',     '/health',        'Unauthenticated health check endpoint'),
    ],
    col_widths=[3.5, 4.5, 8.5],
)
page_break(doc)


# ════════════════════════════════════════════════════
# 7. AUTH FLOW
# ════════════════════════════════════════════════════
section_label(doc, 'Section 7')
heading(doc, '7.  Authentication & Authorization Flow')

sub_heading(doc, 'JWT Token Architecture')
make_table(doc,
    ['Token Type', 'Algorithm', 'Expiry', 'Contains'],
    [
        ('Access Token',  'HS256', '30 minutes (configurable)', 'sub (user/machine ID), username, role, type'),
        ('Refresh Token', 'HS256', '7 days (configurable)',     'sub, type: "refresh"'),
    ],
    col_widths=[3.5, 2.5, 4, 6.5],
)

sub_heading(doc, 'Two Principal Types')
bullet(doc, ('Admin/User accounts — Log in via /v1/auth/login. JWT payload includes role field. '
             'Used by the React admin panel.'))
bullet(doc, ('Machine accounts — Log in via /v1/auth/machine-login. JWT payload includes '
             'type: "machine" and machine_id. Used by the Flutter POS app. '
             'Machine status is updated to "online" on login.'))

sub_heading(doc, 'Authorization Notes')
for b in [
    'All protected endpoints use a get_current_user FastAPI dependency that decodes and validates the Bearer token.',
    'The dependency resolves whether the caller is a User or Machine based on the JWT type claim.',
    'No fine-grained RBAC beyond admin/machine distinction is currently implemented — this is a potential review area.',
    'Logout is client-side only (no server-side token blacklisting).',
    'Inactive users (is_active != "true") are rejected at login.',
    'Machines in maintenance status are rejected at machine-login.',
]:
    bullet(doc, b)

sub_heading(doc, 'Password Storage')
body(doc, ('All passwords (admin users and machine credentials) are hashed using Bcrypt before storage. '
           'Plaintext passwords are never stored.'))

sub_heading(doc, 'Flutter Client Token Handling')
bullet(doc, 'JWT tokens are stored in SharedPreferences on the Android device.')
bullet(doc, 'An HTTP interceptor (Dio _ErrorInterceptor) wraps network errors into typed NetworkException objects.')
bullet(doc, 'Token refresh happens automatically when a 401 is encountered.')

page_break(doc)


# ════════════════════════════════════════════════════
# 8. PAYMENT FLOWS
# ════════════════════════════════════════════════════
section_label(doc, 'Section 8')
heading(doc, '8.  Payment Flows')

sub_heading(doc, 'Flow 1 — Online Payment (UPI via Paytm EDC)')
for s in [
    'Cashier selects services on POS app and initiates payment.',
    'App calls Paytm EDC SDK with amount and merchant credentials.',
    'Paytm SDK handles QR display / card tap / UPI collect internally.',
    'SDK returns a transaction result (success/failure) to the app.',
    'App creates a payment record by calling POST /v1/payments with status, amount, method.',
    'Bill is printed via the POS thermal printer.',
]:
    numbered(doc, s)

sub_heading(doc, 'Flow 2 — Offline Payment (No Network)')
for s in [
    'POS app detects network unavailability (catches NetworkException).',
    'Payment record is enqueued locally in SharedPreferences (JSON list).',
    'A "pending" payment object is returned to the UI so the bill can still be printed.',
    'On the next successful API call, flushSyncQueue() is triggered.',
    'Queued payments are sent to POST /v1/sync/push in bulk.',
    'Server deduplicates by bill_number before storing.',
]:
    numbered(doc, s)

sub_heading(doc, 'Flow 3 — Cash Payment')
for s in [
    'Cashier selects Cash payment method.',
    'App directly calls POST /v1/payments with method: "Cash".',
    'No Paytm SDK interaction. Bill printed immediately.',
]:
    numbered(doc, s)

sub_heading(doc, 'Payment Data Model')
make_table(doc,
    ['Field', 'Type', 'Notes'],
    [
        ('id',          'UUID',      'Auto-generated'),
        ('machine_id',  'UUID (FK)', 'Which POS terminal'),
        ('bill_number', 'String',    'Sequential; used for deduplication in sync'),
        ('amount',      'Decimal',   'Transaction amount in INR'),
        ('method',      'Enum',      'UPI / Card / Cash'),
        ('status',      'Enum',      'success / pending / failed'),
        ('created_at',  'Timestamp', 'UTC; can be set by client for offline payments'),
    ],
    col_widths=[3.5, 3, 10],
)
page_break(doc)


# ════════════════════════════════════════════════════
# 9. DATA SENSITIVITY
# ════════════════════════════════════════════════════
section_label(doc, 'Section 9')
heading(doc, '9.  Data Sensitivity & Classification')

make_table(doc,
    ['Data Type', 'Sensitivity', 'Where Stored', 'Notes'],
    [
        ('Admin credentials',       '🔴 Critical', 'PostgreSQL (bcrypt hashed)',    'Full platform access if compromised'),
        ('Machine credentials',     '🔴 Critical', 'PostgreSQL (bcrypt hashed)',    'Can submit fraudulent payments if compromised'),
        ('JWT Secret Key',          '🔴 Critical', 'Backend environment variable',  'Compromise allows forging any token'),
        ('JWT tokens (access/refresh)', '🔴 Critical', 'Android SharedPreferences', 'No secure storage (Keystore) currently used'),
        ('UPI ID (merchant)',       '🟠 High',     'DB + app memory',              'Admin-configured; rendered in QR code'),
        ('Transaction records',     '🟠 High',     'PostgreSQL',                   'Financial data — amount, method, timestamp'),
        ('Offline payment queue',   '🟠 High',     'Android SharedPreferences (JSON)', 'Unencrypted local storage on device'),
        ('Activity logs',           '🟡 Medium',   'PostgreSQL',                   'Machine activity logs'),
        ('Organisation name, bill prefix', '🟢 Low', 'PostgreSQL / app',           'Non-sensitive configuration'),
    ],
    col_widths=[4, 2.5, 4.5, 5.5],
)

info_box(doc,
    '🔴 Key Risk: JWT Tokens in SharedPreferences',
    ('The Flutter app currently stores JWT tokens in Android SharedPreferences, which is accessible '
     'to root users and backup-extractable on non-hardened devices. The VAPT team should evaluate '
     'whether Android Keystore integration is required for the Paytm POS deployment context.'),
    kind='danger')

page_break(doc)


# ════════════════════════════════════════════════════
# 10. VAPT SCOPE DEFINITION
# ════════════════════════════════════════════════════
section_label(doc, 'Section 10')
heading(doc, '10.  VAPT Scope Definition')

sub_heading(doc, 'In-Scope — API & Backend')
for b in [
    'All REST API endpoints listed in Section 6',
    'Authentication and token management logic',
    'Input validation and SQL injection surface (SQLAlchemy ORM)',
    'Authorization checks (can Machine X access Machine Y\'s payments?)',
    'CORS policy enforcement',
    'Rate limiting (presence/absence)',
    'Error message information disclosure',
    'API documentation exposure (/docs, /redoc) in production',
    'Sync endpoint integrity (can a machine submit payments for another machine?)',
]:
    bullet(doc, b)

sub_heading(doc, 'In-Scope — Mobile Application (Android / Paytm POS)')
for b in [
    'APK static analysis (reverse engineering, hardcoded secrets, API keys)',
    'Data storage analysis (SharedPreferences, SQLite, logs)',
    'Network traffic analysis (certificate pinning, TLS version, MITM)',
    'Improper session handling (token storage, expiry enforcement)',
    'Exported Activities / Broadcast Receivers / Content Providers',
    'Paytm EDC SDK integration (result handling, amount tampering)',
    'Offline queue tampering (injecting/modifying queued payments locally)',
    'Insecure logging (LogCat exposure of sensitive data)',
]:
    bullet(doc, b)

sub_heading(doc, 'In-Scope — Admin Web Panel')
for b in [
    'XSS (Cross-Site Scripting)',
    'CSRF (Cross-Site Request Forgery)',
    'Broken access control (accessing other merchants\' data)',
    'Sensitive data exposure in browser storage (localStorage, sessionStorage, cookies)',
    'Content Security Policy headers',
]:
    bullet(doc, b)

page_break(doc)


# ════════════════════════════════════════════════════
# 11. TESTING FOCUS AREAS
# ════════════════════════════════════════════════════
section_label(doc, 'Section 11')
heading(doc, '11.  Testing Focus Areas & Test Cases')

make_table(doc,
    ['#', 'Category', 'Test Cases', 'Priority'],
    [
        ('1', 'Authentication',
         ('Brute force on /auth/login and /auth/machine-login\n'
          'JWT algorithm confusion (none algorithm, RS256 downgrade)\n'
          'Refresh token reuse after logout\n'
          'Token expiry not enforced client-side\n'
          'Username enumeration via response timing/message difference'),
         '🔴 Critical'),
        ('2', 'Broken Access Control',
         ('Machine A using JWT to access Machine B\'s payments\n'
          'Machine JWT accessing admin-only endpoints\n'
          'IDOR on /payments/{id}, /machines/{id}\n'
          'Horizontal privilege escalation between locations'),
         '🔴 Critical'),
        ('3', 'Payment Integrity',
         ('Amount tampering in API request (POST /payments)\n'
          'Status manipulation (setting status: "success" for a failed txn)\n'
          'Bill number manipulation (deduplication bypass in sync)\n'
          'Submitting payment for a machine_id the caller doesn\'t own\n'
          'Offline queue injection on rooted device'),
         '🔴 Critical'),
        ('4', 'Injection',
         ('SQL injection via query parameters (machine_id, start_date, end_date, etc.)\n'
          'NoSQL injection (not applicable but confirm)\n'
          'XSS in admin panel (machine name, service name, location name)\n'
          'Command injection (not expected but verify)'),
         '🔴 Critical'),
        ('5', 'Sensitive Data Exposure',
         ('JWT secret in environment / source code / logs\n'
          'Merchant UPI ID exposed unnecessarily\n'
          'SharedPreferences readable on non-rooted devices\n'
          'LogCat printing tokens or payment data\n'
          'Stack traces in production API responses (details field)'),
         '🟠 High'),
        ('6', 'API Security',
         ('Missing rate limiting on login endpoints\n'
          '/docs and /redoc accessible in production\n'
          'CORS misconfiguration (wildcard origins)\n'
          'HTTP methods not restricted\n'
          'Response contains excess data (over-exposure)'),
         '🟠 High'),
        ('7', 'Mobile App Static Analysis',
         ('Hardcoded API base URL, secrets, or credentials in APK\n'
          'ProGuard / code obfuscation check\n'
          'Backup flag in AndroidManifest (allowBackup)\n'
          'Debuggable flag in release build\n'
          'Exported components (Activities, Services, Receivers)'),
         '🟠 High'),
        ('8', 'Network Security',
         ('TLS version and cipher suite evaluation\n'
          'Certificate pinning check (if implemented)\n'
          'Man-in-the-middle via proxy (Burp Suite)\n'
          'Cleartext HTTP fallback\n'
          'Network Security Config in Android app'),
         '🟠 High'),
        ('9', 'Business Logic',
         ('Creating payments with negative amounts\n'
          'Creating payments with zero amount\n'
          'Pagination manipulation (negative page, limit=0)\n'
          'Date filter manipulation to extract all historical data\n'
          'Bill counter overflow/manipulation'),
         '🟡 Medium'),
    ],
    col_widths=[0.8, 3.5, 9, 2.5],
)
page_break(doc)


# ════════════════════════════════════════════════════
# 12. OUT OF SCOPE
# ════════════════════════════════════════════════════
section_label(doc, 'Section 12')
heading(doc, '12.  Out-of-Scope Items')

info_box(doc,
    '⚠ Do Not Test',
    ('The following are explicitly out of scope. Testing these without additional authorization '
     'may violate third-party agreements.'),
    kind='warning')

make_table(doc,
    ['Item', 'Reason'],
    [
        ('Paytm EDC SDK internals / Paytm Payment Gateway', 'Third-party system owned and secured by Paytm. Not our infrastructure.'),
        ('Paytm POS device firmware / OS', 'Hardware-level security is Paytm\'s responsibility.'),
        ('Production environment / live data', 'All testing must be on the designated staging environment only.'),
        ('Physical device tampering', 'Out of software security scope.'),
        ('DoS / DDoS attacks', 'Out of scope — may affect availability for other users.'),
        ('Social engineering / phishing', 'Not in scope for this engagement.'),
        ('Cloud infrastructure (host OS, network)', 'Only the application layer is in scope.'),
    ],
    col_widths=[7, 9.5],
)
page_break(doc)


# ════════════════════════════════════════════════════
# 13. TEST ENVIRONMENT
# ════════════════════════════════════════════════════
section_label(doc, 'Section 13')
heading(doc, '13.  Testing Environment Details')

info_box(doc,
    '📋 Environment Provision',
    ('The following details will be provided to the VAPT team by the client team before testing begins. '
     'Blanks to be filled in before sharing this document externally.'),
    kind='info')

make_table(doc,
    ['Item', 'Details'],
    [
        ('Staging API Base URL',       '[ To be provided ]'),
        ('Admin Panel URL (Staging)',  '[ To be provided ]'),
        ('Test Admin Credentials',     '[ To be provided securely ]'),
        ('Test Machine Credentials',   '[ To be provided securely — 2 test machines ]'),
        ('APK Download',               '[ To be provided via secure link ]'),
        ('Swagger / OpenAPI Spec',     'https://[staging-url]/docs'),
        ('Database',                   'Isolated staging PostgreSQL (no production data)'),
        ('Paytm EDC Integration',      'Paytm sandbox credentials will be used for testing'),
    ],
    col_widths=[5.5, 11],
)

sub_heading(doc, 'Test Accounts to be Provided')
make_table(doc,
    ['Account Type', 'Role', 'Purpose'],
    [
        ('Admin User 1', 'Super Admin',  'Full admin access for admin panel testing'),
        ('Machine 1',    'POS Machine',  'Primary test POS account'),
        ('Machine 2',    'POS Machine',  'Cross-machine access control testing'),
    ],
    col_widths=[4, 4, 8.5],
)

sub_heading(doc, 'API Documentation')
body(doc, ('Interactive API documentation is available at /docs (Swagger UI) and /redoc (ReDoc) '
           'on the staging server. These provide full schema definitions for all request/response models.'))

page_break(doc)


# ════════════════════════════════════════════════════
# 14. RULES OF ENGAGEMENT
# ════════════════════════════════════════════════════
section_label(doc, 'Section 14')
heading(doc, '14.  Engagement Rules & Contact')

sub_heading(doc, 'Rules of Engagement')
for b in [
    'All testing must be performed only on the designated staging environment. Production is strictly off-limits.',
    'Any critical vulnerability found (especially involving live payment data or credential exposure) must be reported immediately and not exploited further.',
    'Testers must not exfiltrate or retain any data discovered during testing beyond what is needed for proof-of-concept.',
    'No DoS/DDoS attacks — keep load testing within agreed limits.',
    'All findings must be documented and shared in a structured report with CVSS scoring, reproduction steps, impact, and remediation recommendations.',
    'Testing window: [ To be agreed upon and filled in ]',
]:
    bullet(doc, b)

sub_heading(doc, 'Point of Contact')
make_table(doc,
    ['Role', 'Name', 'Contact'],
    [
        ('Project Lead / Technical Owner', '[ Your Name ]',      '[ Your Email / Phone ]'),
        ('Backend Developer',             '[ Name ]',            '[ Email ]'),
        ('VAPT Engagement Manager',       '[ VAPT Team Lead ]',  '[ Email ]'),
    ],
    col_widths=[5.5, 5, 6],
)

sub_heading(doc, 'Expected Deliverables from VAPT Team')
for b in [
    'Vulnerability Assessment Report (VA Report)',
    'Penetration Testing Report with proof-of-concept for each finding',
    'CVSS v3 score for each vulnerability',
    'Executive Summary suitable for management review',
    'Re-test sign-off after remediation',
]:
    bullet(doc, b)

# Footer line
doc.add_paragraph()
p   = doc.add_paragraph()
run = p.add_run('MIT Billing System — VAPT Scope Document v1.0  ·  Confidential & Restricted  ·  March 2026')
run.font.size  = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
p.paragraph_format.space_before = Pt(20)


# ════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════
out_path = '/Users/nallanaharikrishna/PROJECTS/mit/VAPT_Scope_Document.docx'
doc.save(out_path)
print(f'Saved → {out_path}')
